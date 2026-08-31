"""billing/views.py — PDF and Excel generation for invoices"""
import os
from decimal import Decimal
from django.conf import settings
from django.http import HttpResponse
from django.contrib.admin.views.decorators import staff_member_required
from django.shortcuts import get_object_or_404
from .models import Invoice, PurchaseOrder
from office.models import _amount_to_words

LOGO_PATH = os.path.join(settings.BASE_DIR, 'static', 'visa', 'img', 'logo.png')

CURRENCY_WORDS = {
    'INR': 'Rupees',
    'USD': 'US Dollars',
    'EUR': 'Euros',
}

FRACTION_WORDS = {
    'INR': 'Paise',
    'USD': 'Cents',
    'EUR': 'Cents',
}


def _po_amount_in_words(po):
    """Spells out the exact grand_total, including paise/cents, since PO totals are
    no longer rounded to a whole number. Returns just the words (e.g. 'Fifteen
    Thousand Three Hundred and Sixty Eight Paise Only') — caller prepends the
    currency name."""
    total = po.grand_total or Decimal('0')
    whole = int(total)
    frac = int((total - whole) * 100 + Decimal('0.5'))
    if frac >= 100:
        whole += 1
        frac -= 100

    fraction_word = FRACTION_WORDS.get(po.currency, 'Cents')

    words = _amount_to_words(whole)
    if frac > 0:
        words += f' and {_amount_to_words(frac)} {fraction_word}'
    return words + ' Only'

TERMS = [
    'Goods once sold will not be taken back.',
    'Our responsibility ceases absolutely as soon as the goods have been handed over to carriers.',
]

PO_NOTES = [
    'PO items should be supplied with test certificate and the user manual.',
    'Guarantee/warranty certificate should be provided.',
]

BANK_DETAILS = (
    'Beneficiary name: M/s Virtual Instrumentation &amp; Software Applications Pvt Ltd, '
    'Current Account No: 21260210000991, Bank Name: UCO Bank, Chennai - 600087, '
    'RTGS/IFSC Code: UCBA0002126, MICR Code: 600028027'
)


@staff_member_required
def invoice_pdf(request, pk):
    invoice = get_object_or_404(Invoice, pk=pk)
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        import io
    except ImportError:
        return HttpResponse('reportlab is not installed. Run: pip install reportlab', status=500)

    buffer = io.BytesIO()
    page_w, page_h = A4
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=15*mm, leftMargin=15*mm,
        topMargin=12*mm, bottomMargin=12*mm,
    )
    usable_w = page_w - 30*mm

    def style(name, **kw):
        base = {'fontName': 'Helvetica', 'fontSize': 9, 'leading': 12}
        base.update(kw)
        return ParagraphStyle(name, **base)

    s_company  = style('company', fontName='Helvetica-Bold', fontSize=13, alignment=TA_LEFT, leading=16)
    s_sub      = style('sub', fontSize=8.5, alignment=TA_LEFT, leading=11, textColor=colors.HexColor('#333333'))
    s_title    = style('title', fontName='Helvetica-Bold', fontSize=16, alignment=TA_RIGHT, leading=18)
    s_label    = style('label', fontName='Helvetica-Bold', fontSize=8.5)
    s_value    = style('value', fontSize=8.5)
    s_billto   = style('billto', fontSize=9, leading=12)
    s_head     = style('head', fontName='Helvetica-Bold', fontSize=8.5, textColor=colors.white, alignment=TA_CENTER)
    s_cell     = style('cell', fontSize=8.5, leading=11)
    s_cell_sm  = style('cell_sm', fontSize=7.5, leading=9.5, textColor=colors.HexColor('#555555'))
    s_cell_r   = style('cell_r', fontSize=8.5, leading=11, alignment=TA_RIGHT)
    s_cell_c   = style('cell_c', fontSize=8.5, leading=11, alignment=TA_CENTER)
    s_summary_label = style('summary_label', fontSize=9, alignment=TA_RIGHT)
    s_summary_val   = style('summary_val', fontSize=9, alignment=TA_RIGHT)
    s_grand_label   = style('grand_label', fontName='Helvetica-Bold', fontSize=10.5, alignment=TA_RIGHT)
    s_grand_val     = style('grand_val', fontName='Helvetica-Bold', fontSize=10.5, alignment=TA_RIGHT)
    s_footer   = style('footer', fontSize=7.5, leading=10, textColor=colors.HexColor('#444444'))
    s_center8  = style('center8', fontSize=8, alignment=TA_CENTER)

    elements = []

    # ---------------- Header: logo + company + INVOICE title ----------------
    company_block = Paragraph(
        'Virtual Instrumentation &amp; Software Applications Pvt. Ltd.<br/>'
        '<font size="8">Office &amp; works: Vision Tower, Yogam Garden, 15/16/17, '
        'Brindhavan Nagar, Valasaravakkam, Chennai - 600 087<br/>'
        'Phone: 044 24860722 &nbsp;|&nbsp; www.visapvtltd.co.in &nbsp;|&nbsp; support@visapvtltd.co.in<br/>'
        'GST: 33AABCV2361D1ZT</font>',
        s_company,
    )
    title_block = Paragraph('PROFORMA INVOICE' if invoice.document_type == 'proforma' else 'INVOICE', s_title)

    if os.path.exists(LOGO_PATH):
        logo = Image(LOGO_PATH, width=15*mm, height=15*mm)
        header_table = Table(
            [[logo, company_block, title_block]],
            colWidths=[18*mm, usable_w - 18*mm - 45*mm, 45*mm],
        )
    else:
        header_table = Table(
            [[company_block, title_block]],
            colWidths=[usable_w - 45*mm, 45*mm],
        )
    header_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LINEBELOW', (0, 0), (-1, -1), 1, colors.HexColor('#1D4ED8')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 4*mm))

    # ---------------- Bill To + Invoice meta ----------------
    bill_to_html = f'<b>To:</b><br/>{invoice.client.name}<br/>{invoice.client.address}'.replace('\n', '<br/>')
    if invoice.client.phone:
        bill_to_html += f'<br/>Ph: {invoice.client.phone}'
    bill_to = Paragraph(bill_to_html, s_billto)

    def meta_row(label, value):
        return [Paragraph(f'<b>{label}</b>', s_label), Paragraph(str(value) if value else '', s_value)]

    meta_rows = [
        meta_row('Invoice No.', invoice.invoice_no),
        meta_row('Date', invoice.invoice_date.strftime('%d-%m-%Y')),
        meta_row('Your PO No.', invoice.po_no),
        meta_row('PO Date', invoice.po_date.strftime('%d-%m-%Y') if invoice.po_date else ''),
        meta_row('DC No', invoice.dc_no),
        meta_row('RR/LR/RPP No', invoice.rr_lr_rpp_no),
        meta_row('Buyer GST', invoice.client.gstin),
        meta_row('From / To', f'{invoice.from_place} → {invoice.to_place}'),
    ]
    meta_table = Table(meta_rows, colWidths=[28*mm, 40*mm])
    meta_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
    ]))

    top_block = Table(
        [[bill_to, meta_table]],
        colWidths=[usable_w - 70*mm, 70*mm],
    )
    top_block.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LINEBELOW', (0, 0), (-1, -1), 0.5, colors.HexColor('#cccccc')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(top_block)
    elements.append(Spacer(1, 4*mm))

    # ---------------- Line items table ----------------
    headers = ['S.No', 'Description', 'Qty', 'Unit', 'Unit Price', 'Total Price']
    data = [[Paragraph(h, s_head) for h in headers]]

    for i, item in enumerate(invoice.line_items.all(), start=1):
        sub_lines = []
        if item.model_no:
            sub_lines.append(f'Model: {item.model_no}')
        if item.hsn_code:
            sub_lines.append(f'HSN: {item.hsn_code}')
        if item.serial_no:
            sub_lines.append(f'S.No: {item.serial_no}')
        desc_html = item.description
        if sub_lines:
            desc_html += '<br/>' + '<br/>'.join(f'<font color="#555555" size="7.5">{s}</font>' for s in sub_lines)

        data.append([
            Paragraph(str(i), s_cell_c),
            Paragraph(desc_html, s_cell),
            Paragraph(f'{item.qty:g}', s_cell_c),
            Paragraph(item.unit, s_cell_c),
            Paragraph(f'{item.unit_price:,.2f}', s_cell_r),
            Paragraph(f'{item.amount:,.2f}', s_cell_r),
        ])

    col_widths = [usable_w * w for w in [0.07, 0.43, 0.08, 0.08, 0.15, 0.19]]
    items_table = Table(data, colWidths=col_widths, repeatRows=1)
    items_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1D4ED8')),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CCCCCC')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(items_table)
    elements.append(Spacer(1, 4*mm))

    # ---------------- Tax summary block ----------------
    resolved = invoice.resolved_tax_type()
    summary_rows = [
        ['Sub Total', f'{invoice.subtotal:,.2f}'],
    ]
    if invoice.p_and_f:
        summary_rows.append(['Add: P & F', f'{invoice.p_and_f:,.2f}'])
    summary_rows.append(['Taxable Value', f'{invoice.taxable_value:,.2f}'])

    if resolved == 'igst':
        summary_rows.append([f'IGST {invoice.igst_rate:g}%', f'{invoice.igst_amount:,.2f}'])
    else:
        summary_rows.append([f'CGST {invoice.cgst_rate:g}%', f'{invoice.cgst_amount:,.2f}'])
        summary_rows.append([f'SGST {invoice.sgst_rate:g}%', f'{invoice.sgst_amount:,.2f}'])

    if invoice.insurance:
        summary_rows.append(['Insurance', f'{invoice.insurance:,.2f}'])
    if invoice.less_advance:
        summary_rows.append(['Less: Advance', f'-{invoice.less_advance:,.2f}'])
    summary_rows.append(['Round Off', f'{invoice.round_off:,.2f}'])

    summary_data = [
        [Paragraph(label, s_summary_label), Paragraph(val, s_summary_val)]
        for label, val in summary_rows
    ]
    summary_data.append([
        Paragraph('Grand Total', s_grand_label),
        Paragraph(f'&#8377; {invoice.grand_total:,.2f}', s_grand_val),
    ])

    summary_table = Table(summary_data, colWidths=[45*mm, 40*mm])
    summary_table.setStyle(TableStyle([
        ('LINEABOVE', (0, -1), (-1, -1), 1, colors.HexColor('#1D4ED8')),
        ('TOPPADDING', (0, 0), (-1, -1), 3),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))

    wrapper = Table([[None, summary_table]], colWidths=[usable_w - 85*mm, 85*mm])
    wrapper.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP')]))
    elements.append(wrapper)
    elements.append(Spacer(1, 3*mm))

    # ---------------- Amount in words ----------------
    if invoice.amount_in_words:
        elements.append(Paragraph(f'<b>Rupees:</b> {invoice.amount_in_words}', s_value))
        elements.append(Spacer(1, 4*mm))

    # ---------------- Bank details ----------------
    elements.append(Paragraph(BANK_DETAILS, s_footer))
    elements.append(Spacer(1, 4*mm))

    # ---------------- Terms ----------------
    terms_text = invoice.notes.strip() if invoice.notes.strip() else None
    terms_list = [terms_text] if terms_text else TERMS
    for i, t in enumerate(terms_list, start=1):
        elements.append(Paragraph(f'{i}. {t}', s_footer))
    elements.append(Spacer(1, 10*mm))

    # ---------------- Signature ----------------
    sig_table = Table(
        [[Paragraph('', s_center8), Paragraph('<b>Authorized Signatory</b>', s_center8)]],
        colWidths=[usable_w - 55*mm, 55*mm],
    )
    sig_table.setStyle(TableStyle([
        ('LINEABOVE', (1, 0), (1, 0), 0.5, colors.black),
        ('TOPPADDING', (0, 0), (-1, -1), 20),
        ('ALIGN', (1, 0), (1, 0), 'CENTER'),
    ]))
    elements.append(sig_table)

    doc.build(elements)
    buffer.seek(0)
    doc_label = 'Proforma_Invoice' if invoice.document_type == 'proforma' else 'Invoice'
    filename = f'{doc_label}_{invoice.invoice_no}.pdf'
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'inline; filename="{filename}"'
    return response


@staff_member_required
def invoice_excel(request, pk):
    invoice = get_object_or_404(Invoice, pk=pk)
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter
        from io import BytesIO
    except ImportError:
        return HttpResponse('openpyxl is not installed. Run: pip install openpyxl', status=500)

    wb = Workbook()
    ws = wb.active
    ws.title = f'Invoice {invoice.invoice_no}'

    bold = Font(bold=True)
    header_fill = PatternFill('solid', fgColor='1D4ED8')
    header_font = Font(bold=True, color='FFFFFF')
    thin = Side(style='thin', color='CCCCCC')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal='center', vertical='center')
    wrap = Alignment(wrap_text=True, vertical='top')
    right = Alignment(horizontal='right')

    LAST_COL = 6  # A..F

    # --- Company header ---
    ws.merge_cells('A1:F1')
    ws['A1'] = 'VIRTUAL INSTRUMENTATION & SOFTWARE APPLICATIONS PVT. LTD.'
    ws['A1'].font = Font(bold=True, size=13)
    ws['A1'].alignment = center

    ws.merge_cells('A2:F2')
    ws['A2'] = 'Office & works: Vision Tower, Yogam Garden, 15/16/17, Brindhavan Nagar, Valasaravakkam, Chennai - 600 087'
    ws['A2'].alignment = center
    ws['A2'].font = Font(size=9)

    ws.merge_cells('A3:F3')
    ws['A3'] = 'Phone: 044 24860722  |  www.visapvtltd.co.in  |  support@visapvtltd.co.in  |  GST: 33AABCV2361D1ZT'
    ws['A3'].alignment = center
    ws['A3'].font = Font(size=9)

    ws.merge_cells('A5:F5')
    doc_title = 'PROFORMA INVOICE' if invoice.document_type == 'proforma' else 'INVOICE'
    ws['A5'] = f'{doc_title} — {invoice.invoice_no}'
    ws['A5'].font = Font(bold=True, size=13, color='1D4ED8')
    ws['A5'].alignment = center

    # --- Bill To / meta ---
    row = 7
    ws.cell(row=row, column=1, value='To:').font = bold
    ws.cell(row=row, column=4, value='Invoice No.').font = bold
    ws.cell(row=row, column=5, value=invoice.invoice_no)
    ws.cell(row=row, column=6, value=invoice.invoice_date.strftime('%d-%m-%Y'))
    row += 1
    ws.cell(row=row, column=1, value=invoice.client.name)
    ws.cell(row=row, column=4, value='Your PO No.').font = bold
    ws.cell(row=row, column=5, value=invoice.po_no)
    row += 1
    for line in (invoice.client.address or '').splitlines():
        if line.strip():
            ws.cell(row=row, column=1, value=line.strip())
            row += 1
    if invoice.client.phone:
        ws.cell(row=row, column=1, value=f'Ph: {invoice.client.phone}')
        row += 1
    ws.cell(row=row, column=4, value='DC No').font = bold
    ws.cell(row=row, column=5, value=invoice.dc_no)
    row += 1
    ws.cell(row=row, column=4, value='RR/LR/RPP No').font = bold
    ws.cell(row=row, column=5, value=invoice.rr_lr_rpp_no)
    row += 1
    ws.cell(row=row, column=4, value='Buyer GST').font = bold
    ws.cell(row=row, column=5, value=invoice.client.gstin)
    row += 1
    ws.cell(row=row, column=4, value='From / To').font = bold
    ws.cell(row=row, column=5, value=f'{invoice.from_place} -> {invoice.to_place}')
    row += 2

    # --- Line items ---
    headers = ['S.No', 'Description', 'Qty', 'Unit', 'Unit Price', 'Total Price']
    header_row = row
    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border
    row += 1

    for i, item in enumerate(invoice.line_items.all(), start=1):
        desc = item.description
        extra = []
        if item.model_no:
            extra.append(f'Model: {item.model_no}')
        if item.hsn_code:
            extra.append(f'HSN: {item.hsn_code}')
        if item.serial_no:
            extra.append(f'S.No: {item.serial_no}')
        if extra:
            desc += '\n' + '\n'.join(extra)

        ws.cell(row=row, column=1, value=i).border = border
        ws.cell(row=row, column=1).alignment = center
        c = ws.cell(row=row, column=2, value=desc)
        c.border = border
        c.alignment = wrap
        ws.cell(row=row, column=3, value=float(item.qty)).border = border
        ws.cell(row=row, column=3).alignment = center
        ws.cell(row=row, column=4, value=item.unit).border = border
        ws.cell(row=row, column=4).alignment = center
        p_cell = ws.cell(row=row, column=5, value=float(item.unit_price))
        p_cell.number_format = '#,##0.00'
        p_cell.border = border
        t_cell = ws.cell(row=row, column=6, value=float(item.amount))
        t_cell.number_format = '#,##0.00'
        t_cell.border = border
        row += 1

    row += 1

    # --- Tax summary ---
    resolved = invoice.resolved_tax_type()
    summary_rows = [('Sub Total', invoice.subtotal)]
    if invoice.p_and_f:
        summary_rows.append(('Add: P & F', invoice.p_and_f))
    summary_rows.append(('Taxable Value', invoice.taxable_value))
    if resolved == 'igst':
        summary_rows.append((f'IGST {invoice.igst_rate:g}%', invoice.igst_amount))
    else:
        summary_rows.append((f'CGST {invoice.cgst_rate:g}%', invoice.cgst_amount))
        summary_rows.append((f'SGST {invoice.sgst_rate:g}%', invoice.sgst_amount))
    if invoice.insurance:
        summary_rows.append(('Insurance', invoice.insurance))
    if invoice.less_advance:
        summary_rows.append(('Less: Advance', -invoice.less_advance))
    summary_rows.append(('Round Off', invoice.round_off))

    for label, val in summary_rows:
        ws.cell(row=row, column=5, value=label).font = bold
        ws.cell(row=row, column=5).alignment = right
        v_cell = ws.cell(row=row, column=6, value=float(val))
        v_cell.number_format = '#,##0.00'
        row += 1

    ws.cell(row=row, column=5, value='Grand Total').font = Font(bold=True, size=11, color='1D4ED8')
    ws.cell(row=row, column=5).alignment = right
    gt_cell = ws.cell(row=row, column=6, value=float(invoice.grand_total))
    gt_cell.font = Font(bold=True, size=11, color='1D4ED8')
    gt_cell.number_format = '#,##0.00'
    row += 2

    if invoice.amount_in_words:
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
        ws.cell(row=row, column=1, value=f'Rupees: {invoice.amount_in_words}').font = bold
        row += 2

    # --- Column widths ---
    widths = [8, 42, 8, 10, 14, 16]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    doc_label = 'Proforma_Invoice' if invoice.document_type == 'proforma' else 'Invoice'
    filename = f'{doc_label}_{invoice.invoice_no}.xlsx'
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@staff_member_required
def po_pdf(request, pk):
    po = get_object_or_404(PurchaseOrder, pk=pk)
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, KeepTogether
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
        import io
    except ImportError:
        return HttpResponse('reportlab is not installed. Run: pip install reportlab', status=500)

    buffer = io.BytesIO()
    page_w, page_h = A4
    HEADER_H = 42 * mm   # space reserved at the top of every page for the repeating letterhead

    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=15*mm, leftMargin=15*mm,
        topMargin=HEADER_H, bottomMargin=15*mm,
    )
    usable_w = page_w - 30*mm

    def style(name, **kw):
        base = {'fontName': 'Helvetica', 'fontSize': 9, 'leading': 13}
        base.update(kw)
        return ParagraphStyle(name, **base)

    s_normal  = style('normal', fontSize=9, leading=13)
    s_bold    = style('bold', fontName='Helvetica-Bold', fontSize=9, leading=13)
    s_billto  = style('billto', fontSize=9.5, leading=13)
    s_label   = style('label', fontName='Helvetica-Bold', fontSize=9, leading=13)
    s_value   = style('value', fontSize=9, leading=13)
    s_right   = style('right', fontSize=10, leading=14, alignment=TA_RIGHT)
    s_head    = style('head', fontName='Helvetica-Bold', fontSize=8.5, textColor=colors.white, alignment=TA_CENTER)
    s_cell    = style('cell', fontSize=8.5, leading=11)
    s_cell_c  = style('cell_c', fontSize=8.5, leading=11, alignment=TA_CENTER)
    s_cell_r  = style('cell_r', fontSize=8.5, leading=11, alignment=TA_RIGHT)
    s_center8 = style('center8', fontSize=8, alignment=TA_CENTER)

    # ------------------------------------------------------------------
    # Repeating letterhead — drawn on every page via onPage callback,
    # since Platypus flowables only render once wherever they land.
    # ------------------------------------------------------------------
    def draw_header(canvas, doc_):
        canvas.saveState()
        top_y = page_h - 15*mm
        left_x = 15*mm

        logo_w = 0
        if os.path.exists(LOGO_PATH):
            logo_size = 15*mm
            canvas.drawImage(LOGO_PATH, left_x, top_y - logo_size + 3*mm,
                              width=logo_size, height=logo_size,
                              preserveAspectRatio=True, mask='auto')
            logo_w = logo_size + 4*mm

        text_x = left_x + logo_w
        canvas.setFont('Helvetica-Bold', 15)
        canvas.drawString(text_x, top_y, 'Virtual Instrumentation & Software ')
        canvas.drawString(text_x, top_y - 5.5*mm, 'Applications Pvt. Ltd.')

        canvas.setFont('Helvetica', 8.5)
        canvas.drawString(text_x, top_y - 11.5*mm,
                           'Office & works: Vision Tower, Yogam Garden, 15/16/17, Brindhavan Nagar, Valasaravakkam, Chennai – 600 087')
        canvas.drawString(text_x, top_y - 16*mm,
                           'Ph No: +91 44-24860722, 91-9445350717   Website: www.visapvtltd.net   Email: support@visapvtltd.co.in')

        canvas.setFont('Helvetica-Bold', 18)
        canvas.drawRightString(page_w - 15*mm, top_y, 'PURCHASE ORDER')

        canvas.setStrokeColor(colors.HexColor('#1D4ED8'))
        canvas.setLineWidth(1)
        canvas.line(15*mm, page_h - HEADER_H + 4*mm, page_w - 15*mm, page_h - HEADER_H + 4*mm)
        canvas.restoreState()

    elements = []

    # ---------------- Left: To block (+ Seller GST) | Right: meta panel ----------------
    to_lines = [f'<b>To:</b>', po.supplier.name] + \
               [l.strip() for l in (po.supplier.address or '').splitlines() if l.strip()]
    if po.supplier.phone:
        to_lines.append(f'Ph: {po.supplier.phone}')
    if po.contact_person:
        to_lines.append(f'Kind Attention: {po.contact_person}')
    if po.supplier.gstin:
        to_lines.append(f'<b>Seller GST No:</b> {po.supplier.gstin}')
    bill_to = Paragraph('<br/>'.join(to_lines), s_billto)

    def meta_row(label, value):
        return [Paragraph(f'<b>{label}</b>', s_label), Paragraph(value or '', s_value)]

    from_to = f'{po.from_place} → {po.to_place}' if po.to_place else po.from_place
    meta_rows = [
        meta_row('PO No.', po.po_no),
        meta_row('Date', po.po_date.strftime('%d-%m-%Y')),
        meta_row('Ref No.', po.ref_no),
        meta_row('Ref Date', po.ref_date.strftime('%d-%m-%Y') if po.ref_date else ''),
        meta_row('Currency', po.currency),
        meta_row('Delivery Date', po.delivery_date.strftime('%d-%m-%Y') if po.delivery_date else ''),
        meta_row('From / To', from_to),
        [Paragraph('<b>Payment Terms</b>', s_label), Paragraph(po.payment_terms or '', s_value)],
    ]
    meta_table = Table(meta_rows, colWidths=[28*mm, 42*mm])
    meta_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
    ]))

    top_block = Table([[bill_to, meta_table]], colWidths=[usable_w - 72*mm, 72*mm])
    top_block.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LINEBELOW', (0, 0), (-1, -1), 0.5, colors.HexColor('#cccccc')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
    ]))
    elements.append(top_block)
    elements.append(Spacer(1, 5*mm))

    elements.append(Paragraph('<b>Dear Sir</b>', s_normal))
    elements.append(Spacer(1, 3*mm))
    elements.append(Paragraph('<b>We take pleasure in placing our PO for the following items:</b>', s_normal))
    elements.append(Spacer(1, 3*mm))

    # ---------------- Line items table ----------------
    headers = ['S.No', 'Item', f'Unit Price {po.currency}', 'Qty', f'Total {po.currency}']
    data = [[Paragraph(h, s_head) for h in headers]]
    for i, item in enumerate(po.line_items.all(), start=1):
        data.append([
            Paragraph(str(i), s_cell_c),
            Paragraph(item.description, s_cell),
            Paragraph(f'{item.unit_price:,.2f}', s_cell_r),
            Paragraph(f'{item.qty:g} {item.unit}', s_cell_c),
            Paragraph(f'{item.amount:,.2f}', s_cell_r),
        ])

    resolved = po.resolved_tax_type()
    if resolved == 'cgst_sgst':
        data.append(['', '', '', Paragraph('CGST', s_bold), Paragraph(f'{po.cgst_amount:,.2f}', s_cell_r)])
        data.append(['', '', '', Paragraph('SGST', s_bold), Paragraph(f'{po.sgst_amount:,.2f}', s_cell_r)])
    elif resolved == 'igst':
        data.append(['', '', '', Paragraph('IGST', s_bold), Paragraph(f'{po.igst_amount:,.2f}', s_cell_r)])
    data.append(['', '', '', Paragraph('<b>Total</b>', s_bold), Paragraph(f'<b>{po.grand_total:,.2f}</b>', s_cell_r)])

    col_widths = [usable_w * w for w in [0.07, 0.40, 0.18, 0.15, 0.20]]
    items_table = Table(data, colWidths=col_widths, repeatRows=1)
    items_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1D4ED8')),
        ('GRID', (0, 0), (-1, -1), 0.4, colors.HexColor('#CCCCCC')),
        ('LINEABOVE', (3, -1), (-1, -1), 0.75, colors.black),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
    ]))
    elements.append(items_table)
    elements.append(Spacer(1, 4*mm))

    # ---------------- Amount in words ----------------
    words = _po_amount_in_words(po) if po.grand_total else ''
    currency_word = CURRENCY_WORDS.get(po.currency, po.currency)
    elements.append(Paragraph(f'<b>{currency_word}:</b> {words}', s_normal))
    elements.append(Spacer(1, 6*mm))

    # ---------------- Notes + closing + signature — kept together as one unbreakable block ----------------
    notes_list = [n.strip() for n in po.notes.strip().splitlines() if n.strip()] if po.notes.strip() else PO_NOTES
    closing_block = []
    for i, n in enumerate(notes_list, start=1):
        closing_block.append(Paragraph(f'{i}. {n}', s_normal))
    closing_block.append(Spacer(1, 6*mm))
    closing_block.append(Paragraph('Thanking you', s_normal))
    closing_block.append(Spacer(1, 14*mm))

    sig_table = Table(
        [[Paragraph('', s_center8), Paragraph('<b>Authorized Signatory</b>', s_center8)]],
        colWidths=[usable_w - 55*mm, 55*mm],
    )
    sig_table.setStyle(TableStyle([
        ('LINEABOVE', (1, 0), (1, 0), 0.5, colors.black),
        ('TOPPADDING', (0, 0), (-1, -1), 20),
        ('ALIGN', (1, 0), (1, 0), 'CENTER'),
    ]))
    closing_block.append(sig_table)

    elements.append(KeepTogether(closing_block))

    doc.build(elements, onFirstPage=draw_header, onLaterPages=draw_header)
    buffer.seek(0)
    filename = f'PO_{po.po_no.replace("/", "-")}.pdf'
    response = HttpResponse(buffer, content_type='application/pdf')
    response['Content-Disposition'] = f'inline; filename="{filename}"'
    return response



@staff_member_required
def po_excel(request, pk):
    po = get_object_or_404(PurchaseOrder, pk=pk)
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
        from openpyxl.utils import get_column_letter
        from io import BytesIO
    except ImportError:
        return HttpResponse('openpyxl is not installed. Run: pip install openpyxl', status=500)

    wb = Workbook()
    ws = wb.active
    ws.title = f'PO {po.po_no.replace("/", "-")}'

    bold = Font(bold=True)
    header_fill = PatternFill('solid', fgColor='1D4ED8')
    header_font = Font(bold=True, color='FFFFFF')
    thin = Side(style='thin', color='CCCCCC')
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal='center', vertical='center')
    wrap = Alignment(wrap_text=True, vertical='top')

    ws.merge_cells('A1:F1')
    ws['A1'] = 'VIRTUAL INSTRUMENTATION & SOFTWARE APPLICATIONS PVT LTD.'
    ws['A1'].font = Font(bold=True, size=13)
    ws['A1'].alignment = center

    ws.merge_cells('A2:F2')
    ws['A2'] = 'PURCHASE ORDER'
    ws['A2'].font = Font(bold=True, size=13, color='1D4ED8')
    ws['A2'].alignment = center

    # --- Left: To block (with Seller GST) | Right: meta panel ---
    row = 4
    ws.cell(row=row, column=1, value='To:').font = bold
    ws.cell(row=row, column=4, value='PO No.').font = bold
    ws.cell(row=row, column=5, value=po.po_no)
    row += 1
    ws.cell(row=row, column=1, value=po.supplier.name)
    ws.cell(row=row, column=4, value='Date').font = bold
    ws.cell(row=row, column=5, value=po.po_date.strftime('%d-%m-%Y'))
    row += 1
    for line in (po.supplier.address or '').splitlines():
        if line.strip():
            ws.cell(row=row, column=1, value=line.strip())
            row += 1
    if po.supplier.phone:
        ws.cell(row=row, column=1, value=f'Ph: {po.supplier.phone}')
        row += 1
    if po.contact_person:
        ws.cell(row=row, column=1, value=f'Kind Attention: {po.contact_person}')
        row += 1
    if po.supplier.gstin:
        ws.cell(row=row, column=1, value=f'Seller GST No: {po.supplier.gstin}').font = bold
        row += 1

    # Continue meta panel on the right, starting from row 6 regardless of how tall the To block got
    meta_row_num = 6
    from_to = f'{po.from_place} -> {po.to_place}' if po.to_place else po.from_place
    meta_pairs = [
        ('Ref No.', po.ref_no),
        ('Ref Date', po.ref_date.strftime('%d-%m-%Y') if po.ref_date else ''),
        ('Currency', po.currency),
        ('Delivery Date', po.delivery_date.strftime('%d-%m-%Y') if po.delivery_date else ''),
        ('From / To', from_to),
        ('Payment Terms', po.payment_terms),
    ]
    for label, val in meta_pairs:
        ws.cell(row=meta_row_num, column=4, value=label).font = bold
        ws.cell(row=meta_row_num, column=5, value=val)
        meta_row_num += 1

    row = max(row, meta_row_num) + 1

    headers = ['S.No', 'Item', f'Unit Price {po.currency}', 'Qty', 'Unit', f'Total {po.currency}']
    header_row = row
    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = border
    row += 1

    for i, item in enumerate(po.line_items.all(), start=1):
        ws.cell(row=row, column=1, value=i).border = border
        ws.cell(row=row, column=1).alignment = center
        c = ws.cell(row=row, column=2, value=item.description)
        c.border = border
        c.alignment = wrap
        p_cell = ws.cell(row=row, column=3, value=float(item.unit_price))
        p_cell.number_format = '#,##0.00'
        p_cell.border = border
        ws.cell(row=row, column=4, value=float(item.qty)).border = border
        ws.cell(row=row, column=4).alignment = center
        ws.cell(row=row, column=5, value=item.unit).border = border
        ws.cell(row=row, column=5).alignment = center
        t_cell = ws.cell(row=row, column=6, value=float(item.amount))
        t_cell.number_format = '#,##0.00'
        t_cell.border = border
        row += 1

    row += 1
    resolved = po.resolved_tax_type()
    if resolved == 'cgst_sgst':
        ws.cell(row=row, column=5, value='CGST').font = bold
        ws.cell(row=row, column=6, value=float(po.cgst_amount)).number_format = '#,##0.00'
        row += 1
        ws.cell(row=row, column=5, value='SGST').font = bold
        ws.cell(row=row, column=6, value=float(po.sgst_amount)).number_format = '#,##0.00'
        row += 1
    elif resolved == 'igst':
        ws.cell(row=row, column=5, value='IGST').font = bold
        ws.cell(row=row, column=6, value=float(po.igst_amount)).number_format = '#,##0.00'
        row += 1

    ws.cell(row=row, column=5, value='Total').font = Font(bold=True, size=11, color='1D4ED8')
    total_cell = ws.cell(row=row, column=6, value=float(po.grand_total))
    total_cell.font = Font(bold=True, size=11, color='1D4ED8')
    total_cell.number_format = '#,##0.00'
    row += 2

    words = _po_amount_in_words(po) if po.grand_total else ''
    currency_word = CURRENCY_WORDS.get(po.currency, po.currency)
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
    ws.cell(row=row, column=1, value=f'{currency_word}: {words}').font = bold
    row += 2

    notes_list = [n.strip() for n in po.notes.strip().splitlines() if n.strip()] if po.notes.strip() else PO_NOTES
    for i, n in enumerate(notes_list, start=1):
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
        ws.cell(row=row, column=1, value=f'{i}. {n}')
        row += 1
    row += 1

    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=6)
    ws.cell(row=row, column=1, value='Thanking you')
    row += 3
    ws.merge_cells(start_row=row, start_column=5, end_row=row, end_column=6)
    ws.cell(row=row, column=5, value='Authorized Signatory').font = bold
    ws.cell(row=row, column=5).alignment = center

    widths = [8, 44, 16, 10, 10, 16]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    filename = f'PO_{po.po_no.replace("/", "-")}.xlsx'
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@staff_member_required
def po_docx(request, pk):
    po = get_object_or_404(PurchaseOrder, pk=pk)
    try:
        from docx import Document
        from docx.shared import Pt, Mm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from io import BytesIO
    except ImportError:
        return HttpResponse('python-docx is not installed. Run: pip install python-docx', status=500)

    document = Document()
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)

    # ---------------- Header (a real Word header — repeats on every page automatically) ----------------
    header = section.header
    header.is_linked_to_previous = False
    header.paragraphs[0].text = ''

    header_table = header.add_table(rows=1, cols=2, width=Mm(180))
    header_table.autofit = False
    header_table.columns[0].width = Mm(122)
    header_table.columns[1].width = Mm(58)

    left_cell = header_table.cell(0, 0)
    p = left_cell.paragraphs[0]
    run = p.add_run('Virtual Instrumentation & Software Applications Pvt Ltd.')
    run.bold = True
    run.font.size = Pt(13)

    p2 = left_cell.add_paragraph()
    run2 = p2.add_run(
        'Office & works: Vision Tower, Yogam Garden, 15/16/17, Brindhavan Nagar, '
        'Valasaravakkam, Chennai - 600 087'
    )
    run2.font.size = Pt(8.5)

    p3 = left_cell.add_paragraph()
    run3 = p3.add_run(
        'Ph No: +91 44-24860722, 91-9445350717   Website: www.visapvtltd.net   '
        'Email: support@visapvtltd.co.in'
    )
    run3.font.size = Pt(8.5)

    right_cell = header_table.cell(0, 1)
    right_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run4 = right_cell.paragraphs[0].add_run('PURCHASE ORDER')
    run4.bold = True
    run4.font.size = Pt(15)

    hr_p = header.add_paragraph()
    pPr = hr_p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1D4ED8')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ---------------- Body: To block (+ Seller GST) | meta panel ----------------
    top_table = document.add_table(rows=1, cols=2)
    top_table.autofit = False
    top_table.columns[0].width = Mm(108)
    top_table.columns[1].width = Mm(72)

    to_cell = top_table.cell(0, 0)
    to_run = to_cell.paragraphs[0].add_run('To:')
    to_run.bold = True
    to_cell.add_paragraph(po.supplier.name)
    for line in (po.supplier.address or '').splitlines():
        if line.strip():
            to_cell.add_paragraph(line.strip())
    if po.supplier.phone:
        to_cell.add_paragraph(f'Ph: {po.supplier.phone}')
    if po.contact_person:
        to_cell.add_paragraph(f'Kind Attention: {po.contact_person}')
    if po.supplier.gstin:
        gst_p = to_cell.add_paragraph()
        gst_run = gst_p.add_run(f'Seller GST No: {po.supplier.gstin}')
        gst_run.bold = True

    meta_cell = top_table.cell(0, 1)
    from_to = f'{po.from_place} \u2192 {po.to_place}' if po.to_place else po.from_place
    meta_pairs = [
        ('PO No.', po.po_no),
        ('Date', po.po_date.strftime('%d-%m-%Y')),
        ('Ref No.', po.ref_no or ''),
        ('Ref Date', po.ref_date.strftime('%d-%m-%Y') if po.ref_date else ''),
        ('Currency', po.currency),
        ('Delivery Date', po.delivery_date.strftime('%d-%m-%Y') if po.delivery_date else ''),
        ('From / To', from_to),
        ('Payment Terms', po.payment_terms or ''),
    ]
    first = True
    for label, val in meta_pairs:
        p = meta_cell.paragraphs[0] if first else meta_cell.add_paragraph()
        first = False
        r1 = p.add_run(f'{label}: ')
        r1.bold = True
        r1.font.size = Pt(9)
        r2 = p.add_run(val)
        r2.font.size = Pt(9)

    document.add_paragraph()

    dear_p = document.add_paragraph()
    dear_run = dear_p.add_run('Dear Sir')
    dear_run.bold = True

    intro_p = document.add_paragraph()
    intro_run = intro_p.add_run('We take pleasure in placing our PO for the following items:')
    intro_run.bold = True

    # ---------------- Line items table ----------------
    resolved = po.resolved_tax_type()
    items = list(po.line_items.all())
    extra_rows = 0
    if resolved == 'cgst_sgst':
        extra_rows = 2
    elif resolved == 'igst':
        extra_rows = 1
    extra_rows += 1  # Total row

    table = document.add_table(rows=1 + len(items) + extra_rows, cols=5)
    table.style = 'Table Grid'
    col_widths = [Mm(12), Mm(72), Mm(30), Mm(24), Mm(32)]
    for row_cells in table.rows:
        for idx, cell in enumerate(row_cells.cells):
            cell.width = col_widths[idx]

    hdr_cells = table.rows[0].cells
    headers = ['S.No', 'Item', f'Unit Price {po.currency}', 'Qty', f'Total {po.currency}']
    for idx, htext in enumerate(headers):
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(htext)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1D4ED8')
        hdr_cells[idx]._tc.get_or_add_tcPr().append(shd)

    r_idx = 1
    for i, item in enumerate(items, start=1):
        cells = table.rows[r_idx].cells
        cells[0].text = str(i)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[1].text = item.description
        cells[2].text = f'{item.unit_price:,.2f}'
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[3].text = f'{item.qty:g} {item.unit}'
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[4].text = f'{item.amount:,.2f}'
        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_idx += 1

    def _tax_row(label, amount, bold_row=False):
        nonlocal r_idx
        cells = table.rows[r_idx].cells
        run = cells[3].paragraphs[0].add_run(label)
        run.bold = True
        run2 = cells[4].paragraphs[0].add_run(f'{amount:,.2f}')
        run2.bold = bold_row
        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_idx += 1

    if resolved == 'cgst_sgst':
        _tax_row('CGST', po.cgst_amount)
        _tax_row('SGST', po.sgst_amount)
    elif resolved == 'igst':
        _tax_row('IGST', po.igst_amount)
    _tax_row('Total', po.grand_total, bold_row=True)

    document.add_paragraph()

    # ---------------- Amount in words ----------------
    words = _po_amount_in_words(po) if po.grand_total else ''
    currency_word = CURRENCY_WORDS.get(po.currency, po.currency)
    words_p = document.add_paragraph()
    words_r1 = words_p.add_run(f'{currency_word}: ')
    words_r1.bold = True
    words_p.add_run(words)

    document.add_paragraph()

    # ---------------- Notes + Thanking you + signature — kept together as one unbreakable block ----------------
    notes_list = [n.strip() for n in po.notes.strip().splitlines() if n.strip()] if po.notes.strip() else PO_NOTES
    for i, n in enumerate(notes_list, start=1):
        note_p = document.add_paragraph(f'{i}. {n}')
        note_p.paragraph_format.keep_with_next = True

    thanks_p = document.add_paragraph('Thanking you')
    thanks_p.paragraph_format.keep_with_next = True

    spacer_p = document.add_paragraph()
    spacer_p.paragraph_format.keep_with_next = True

    # ---------------- Signature ----------------
    sig_table = document.add_table(rows=1, cols=2)
    sig_table.columns[0].width = Mm(125)
    sig_table.columns[1].width = Mm(55)
    sig_cell = sig_table.cell(0, 1)

    line_p = sig_cell.paragraphs[0]
    pPr2 = line_p._p.get_or_add_pPr()
    pBdr2 = OxmlElement('w:pBdr')
    bottom2 = OxmlElement('w:bottom')
    bottom2.set(qn('w:val'), 'single')
    bottom2.set(qn('w:sz'), '6')
    bottom2.set(qn('w:space'), '1')
    bottom2.set(qn('w:color'), '000000')
    pBdr2.append(bottom2)
    pPr2.append(pBdr2)

    label_p = sig_cell.add_paragraph()
    label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_run = label_p.add_run('Authorized Signatory')
    sig_run.bold = True

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    filename = f'PO_{po.po_no.replace("/", "-")}.docx'
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response